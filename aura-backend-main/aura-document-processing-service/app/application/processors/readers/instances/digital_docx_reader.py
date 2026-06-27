import logging
from pathlib import Path
from typing import Optional
from docx import Document
from docx.document import Document as DocxDocument

from app.application.processors.readers.exceptions.reader_exception import (
    DigitalDOCXReadException,
    DOCXHasNoExtractableTextException,
    ReaderFileNotFoundException,
)
from app.application.processors.readers.instances.base_reader import BaseReader
from app.application.processors.readers.reader_settings import ReaderSettings

logger = logging.getLogger(__name__)

class DigitalDOCXReader(BaseReader):
    def __init__(
            self,
            reader_settings: Optional[ReaderSettings] = None
    ) -> None:
        self._settings = reader_settings or ReaderSettings()

        logger.info("The digital DOCX reader was initialized successfully.")

    def can_handle(
            self,
            file_path: Path
    ) -> bool:
        if file_path.suffix.lower() != ".docx":
            return False

        try:
            doc = Document(str(file_path))
            if any(p.text and p.text.strip() for p in doc.paragraphs[:20]):
                return True
            if doc.tables:
                return any(
                    cell.text and cell.text.strip()
                    for row in doc.tables[0].rows[:5]
                    for cell in row.cells
                )
            return False

        except Exception as e:
            logger.debug(
                "An error occurred while checking whether the DOCX can be handled.",
                extra={
                    "file_name": file_path.name,
                    "exception_type": type(e).__name__
                }
            )
            return False

    def read(
            self,
            file_path: Path
    ) -> str:
        self._validate_file_exists(file_path)

        logger.info(
            "Reading a digital DOCX file.",
            extra={
                "file_name": file_path.name
            }
        )

        try:
            doc = Document(str(file_path))
            text_parts = self._extract_text(doc)

            if not text_parts:
                raise DOCXHasNoExtractableTextException("The DOCX file does not contain extractable text content.")

            logger.info(
                "The digital DOCX was read successfully.",
                extra={
                    "file_name": file_path.name,
                    "parts": len(text_parts)
                }
            )

            return "\n\n".join(text_parts)

        except (
                ReaderFileNotFoundException,
                DOCXHasNoExtractableTextException
        ):
            raise
        except Exception as e:
            logger.exception(
                "An error occurred while reading the digital DOCX.",
                extra={
                    "file_name": file_path.name
                }
            )
            raise DigitalDOCXReadException("An unexpected error occurred while reading the digital DOCX file.") from e

    def _extract_text(
            self,
            doc: DocxDocument
    ) -> list[str]:
        text_parts: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                seen: set[int] = set()
                row_cells: list[str] = []
                for cell in row.cells:
                    if id(cell) in seen:
                        continue
                    seen.add(id(cell))
                    if cell.text and cell.text.strip():
                        row_cells.append(cell.text.strip())
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        return text_parts
